import os
from flask import Flask, request, render_template, url_for, redirect
from docx import Document
import re
import os
import time

app = Flask(__name__)

file_uploads = []
msg = 'hello'
file_link = ''
path = 'uploads/'


def txt2doc(path,title):
    global file_link
    document = Document()
    document.add_heading(title, 0)
    myfile = open(path + title, encoding='utf-8').read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile)  # remove all non-XML-compatible characters
    document.add_paragraph(myfile)
    document.save(path + title + '.docx')
    file_link = path + title + '.docx'

def get_file_link(file_link):
    if file_link:
        return file_link

@app.route("/")
def fileFrontPage():
    global msg, file_link
    if get_file_link(file_link):
        link = file_link
    else:
        link = ''
    return render_template('index.html',message =msg, file=link)

@app.route("/handleUpload", methods=['POST'])
def handleFileUpload():
    if 'photo' in request.files:
        photo = request.files['photo']
        if photo.filename != '':
            file_uploads.append(photo.filename)
            photo.save(os.path.join('uploads', photo.filename))
            msg_text = 'Upload Succesfully'
            if file_uploads:
                from tidf import TDIF
                print(file_uploads)
                tid = TDIF(file_uploads)
                tid.text_doc()
                tid.remove_puncts()
                print(tid.caltf())
                print(tid.idf())
                print(tid.tfidf())
                (tid.frequency_check())
                (tid.freq_answer())
                (tid.var_calc())
                (tid.ent_calc())
                (tid.avg_tfidf())
               # print(tid.new_variance)
                #print(tid.new_entrop)
                print(tid.normal_freq)
                #print(tid.new_tfidf)
                #print(tid.new_frequency)
                a = set(tid.new_variance).intersection(tid.new_entrop)
                b =(a.intersection(tid.new_frequency))
                c = (b.intersection(tid.new_tfidf))
                global msg
                msg = msg_text
                print(c)
                filename_os = "originalstopwords.txt"
                with open(path+filename_os, mode='w', encoding='utf-8') as text_file:
                    for stopwords in c:
                        text_file.write("{}\n".format(stopwords))
                txt2doc(path, filename_os)
            #flash(msg_text)
    return redirect(url_for('fileFrontPage'))

if __name__ == '__main__':
    port = 5001
    app.run(host='0.0.0.0', port=port)
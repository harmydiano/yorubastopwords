"""
Enhanced Flask Application for Yoruba Stopwords Generator

This is the main web application that provides a user-friendly interface
for generating Yoruba stopwords from uploaded documents.
"""

import os
import logging
from pathlib import Path
from typing import Optional
from datetime import datetime

from flask import Flask, request, render_template, url_for, redirect, jsonify, send_file, flash
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from docx import Document

from config import config, Config
from tfidf_analyzer import TFIDFAnalyzer

# Initialize Flask app
app = Flask(__name__)

# Load configuration
env = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[env])
Config.init_app(app)

# Setup logging
logging.basicConfig(
    level=app.config['LOG_LEVEL'],
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(app.config['LOG_FILE']),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Global state
uploaded_files = []
processing_result = {
    'status': None,
    'message': '',
    'stopwords_count': 0,
    'file_links': {},
    'statistics': {}
}


def allowed_file(filename: str) -> bool:
    """
    Check if file extension is allowed.

    Args:
        filename: Name of the file to check

    Returns:
        True if file extension is allowed, False otherwise
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def validate_uploaded_file(file) -> tuple[bool, Optional[str]]:
    """
    Validate uploaded file.

    Args:
        file: Uploaded file object

    Returns:
        Tuple of (is_valid, error_message)
    """
    if not file:
        return False, "No file provided"

    if file.filename == '':
        return False, "No file selected"

    if not allowed_file(file.filename):
        return False, f"Invalid file type. Allowed types: {', '.join(app.config['ALLOWED_EXTENSIONS'])}"

    # Check file size (Flask already handles this with MAX_CONTENT_LENGTH)
    # Additional validation can be added here

    return True, None


def create_docx_from_text(text_path: Path, title: str) -> Path:
    """
    Convert text file to .docx format.

    Args:
        text_path: Path to the text file
        title: Title for the document

    Returns:
        Path to the created .docx file
    """
    try:
        document = Document()
        document.add_heading(title, 0)

        # Read text file
        with open(text_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Add content to document
        document.add_paragraph(content)

        # Save document
        docx_path = text_path.with_suffix('.docx')
        document.save(str(docx_path))

        logger.info(f"Created DOCX file: {docx_path}")
        return docx_path

    except Exception as e:
        logger.error(f"Error creating DOCX file: {e}")
        raise


@app.route("/")
def index():
    """Render the main page."""
    return render_template(
        'index_new.html',
        message=processing_result.get('message', ''),
        status=processing_result.get('status'),
        file_links=processing_result.get('file_links', {}),
        statistics=processing_result.get('statistics', {}),
        stopwords_count=processing_result.get('stopwords_count', 0)
    )


@app.route("/upload", methods=['POST'])
def handle_upload():
    """
    Handle file upload and process stopwords extraction.
    """
    global uploaded_files, processing_result

    # Reset state
    uploaded_files = []
    processing_result = {
        'status': 'error',
        'message': '',
        'stopwords_count': 0,
        'file_links': {},
        'statistics': {}
    }

    try:
        # Check if files are in request
        if 'files[]' not in request.files:
            processing_result['message'] = "No files uploaded"
            flash("No files uploaded", "error")
            return redirect(url_for('index'))

        files = request.files.getlist('files[]')

        if not files or all(f.filename == '' for f in files):
            processing_result['message'] = "No files selected"
            flash("No files selected", "error")
            return redirect(url_for('index'))

        # Validate and save files
        for file in files:
            is_valid, error_msg = validate_uploaded_file(file)

            if not is_valid:
                processing_result['message'] = error_msg
                flash(error_msg, "error")
                return redirect(url_for('index'))

            # Secure the filename
            filename = secure_filename(file.filename)

            # Add timestamp to avoid conflicts
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{timestamp}{ext}"

            # Save file
            filepath = Path(app.config['UPLOAD_FOLDER']) / filename
            file.save(str(filepath))

            uploaded_files.append(filename)
            logger.info(f"Uploaded file: {filename}")

        # Process files
        logger.info(f"Processing {len(uploaded_files)} files")

        # Create analyzer
        analyzer = TFIDFAnalyzer(
            filenames=uploaded_files,
            min_word_length=app.config.get('MIN_WORD_LENGTH', 2),
            max_word_length=app.config.get('MAX_WORD_LENGTH', 5),
            min_frequency=app.config.get('MIN_FREQUENCY', 5)
        )

        # Extract stopwords
        stopwords = analyzer.process()

        # Get statistics
        statistics = analyzer.get_statistics()

        # Save results in multiple formats
        output_base = Path(app.config['OUTPUT_FOLDER']) / 'yoruba_stopwords'
        saved_files = analyzer.save_stopwords(
            output_base,
            formats=['txt', 'json', 'csv']
        )

        # Create DOCX version
        txt_path = saved_files.get('txt')
        if txt_path:
            docx_path = create_docx_from_text(txt_path, 'Yoruba Stopwords')
            saved_files['docx'] = docx_path

        # Update processing result
        processing_result.update({
            'status': 'success',
            'message': f'Successfully extracted {len(stopwords)} stopwords!',
            'stopwords_count': len(stopwords),
            'file_links': {
                fmt: str(path.relative_to(Path(app.config['OUTPUT_FOLDER']).parent))
                for fmt, path in saved_files.items()
            },
            'statistics': statistics
        })

        flash(f'Successfully extracted {len(stopwords)} stopwords!', 'success')
        logger.info(f"Successfully processed files. Found {len(stopwords)} stopwords")

    except RequestEntityTooLarge:
        processing_result['message'] = "File too large. Maximum size is 16MB"
        flash("File too large. Maximum size is 16MB", "error")
        logger.error("File too large")

    except Exception as e:
        processing_result['message'] = f"Error processing files: {str(e)}"
        flash(f"Error processing files: {str(e)}", "error")
        logger.error(f"Error processing files: {e}", exc_info=True)

    return redirect(url_for('index'))


@app.route("/download/<path:filename>")
def download_file(filename):
    """
    Download a generated file.

    Args:
        filename: Name of the file to download
    """
    try:
        file_path = Path(app.config['OUTPUT_FOLDER']) / filename

        if not file_path.exists():
            flash("File not found", "error")
            return redirect(url_for('index'))

        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        flash(f"Error downloading file: {str(e)}", "error")
        return redirect(url_for('index'))


@app.route("/api/stopwords", methods=['GET'])
def api_get_stopwords():
    """
    API endpoint to get stopwords in JSON format.

    Returns:
        JSON response with stopwords and statistics
    """
    try:
        stopwords_file = Path(app.config['OUTPUT_FOLDER']) / 'yoruba_stopwords.json'

        if not stopwords_file.exists():
            return jsonify({
                'error': 'No stopwords generated yet',
                'status': 'not_found'
            }), 404

        with open(stopwords_file, 'r', encoding='utf-8') as f:
            import json
            data = json.load(f)

        return jsonify({
            'status': 'success',
            'data': data
        })

    except Exception as e:
        logger.error(f"API error: {e}")
        return jsonify({
            'error': str(e),
            'status': 'error'
        }), 500


@app.route("/api/health", methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    })


@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors."""
    return render_template('404.html'), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors."""
    logger.error(f"Internal error: {error}")
    return render_template('500.html'), 500


@app.errorhandler(RequestEntityTooLarge)
def handle_large_file(error):
    """Handle file too large errors."""
    flash("File too large. Maximum size is 16MB", "error")
    return redirect(url_for('index'))


# CLI commands
@app.cli.command()
def init_db():
    """Initialize the application."""
    logger.info("Initializing application")
    Config.init_app(app)
    logger.info("Application initialized")


if __name__ == '__main__':
    # Ensure directories exist
    Config.init_app(app)

    # Run application
    app.run(
        host=app.config['HOST'],
        port=app.config['PORT'],
        debug=app.config['DEBUG']
    )
